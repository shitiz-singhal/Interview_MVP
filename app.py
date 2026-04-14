import asyncio

# FORCED FIX FOR PYTHON 3.14 + STREAMLIT CLOUD
try:
    asyncio.get_event_loop()
except RuntimeError:
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)

import streamlit as st
from audio_recorder_streamlit import audio_recorder
import assemblyai as aai
import os 
from anthropic import Anthropic
from docx import Document

# --- SETUP THE FOLDER ---
SAVE_FOLDER = "Evaluations_History"
os.makedirs(SAVE_FOLDER, exist_ok=True)

# Pull keys securely from Streamlit Secrets
aai_api_key = st.secrets.get("AAI_API_KEY", "")
anthropic_api_key = st.secrets.get("ANTHROPIC_API_KEY", "")

st.title("Interview Evaluator MVP")

candidate_name = st.text_input("Candidate Name:")
evaluator_names = st.text_input("Evaluator Names (comma-separated, e.g., Jane, John):")
total_speakers = st.number_input("Total number of people speaking (Candidate + Evaluators):", min_value=2, max_value=10, value=3)

st.divider()

# --- THE NEW TAB STRUCTURE ---
tab1, tab2, tab3, tab4 = st.tabs(["Upload Audio File", "Record Live Audio", "Upload Text Transcript", "Dashboard"])

audio_bytes = None 
transcript_text = None 

# --- TAB 1: UPLOADING AUDIO ---
with tab1:
    uploaded_file = st.file_uploader("Choose an audio file", type=["wav", "mp3", "m4a", "mpeg"])
    if uploaded_file is not None:
        audio_bytes = uploaded_file.read() 
        st.audio(audio_bytes) 
        st.success("Audio file uploaded successfully!")

# --- TAB 2: RECORDING AUDIO ---
with tab2:
    recorded_audio = audio_recorder(text="Record", recording_color="#e81e1e", neutral_color="#6aa36f", pause_threshold=300.0)
    if recorded_audio:
        audio_bytes = recorded_audio 
        st.audio(audio_bytes, format="audio/wav") 
        st.success("Recording captured successfully!")

# --- TAB 3: UPLOADING TEXT TRANSCRIPT ---
with tab3:
    uploaded_text = st.file_uploader("Upload a text transcript (.txt)", type=["txt"])
    if uploaded_text is not None:
        # Read the file and decode it into a standard Python string
        transcript_text = uploaded_text.read().decode("utf-8")
        st.success("Text transcript uploaded successfully!")
        with st.expander("👀 View Uploaded Transcript"):
            st.text(transcript_text)

# --- TAB 4: THE DASHBOARD ---
with tab4:
    st.subheader("Your Past Evaluations")
    
    # CLEAR DASHBOARD BUTTON
    colA, colB = st.columns([3, 1])
    with colB:
        if st.button("🗑️ Clear Dashboard"):
            for file_name in os.listdir(SAVE_FOLDER):
                os.remove(os.path.join(SAVE_FOLDER, file_name))
            st.success("Dashboard cleared!")
            st.rerun()

    saved_files = os.listdir(SAVE_FOLDER)
    
    if len(saved_files) == 0:
        st.info("No documents completed yet.")
    else:
        for file_name in saved_files:
            with st.container():
                col1, col2 = st.columns([3, 1]) 
                col1.write(f"📄 **{file_name}**")
                
                file_path = os.path.join(SAVE_FOLDER, file_name)
                with open(file_path, "rb") as file:
                    col2.download_button(label="Download", data=file, file_name=file_name, key=file_name)
                st.divider()

# --- SUBMIT BUTTON & AI PROCESSING ---
if (audio_bytes or transcript_text) and candidate_name and evaluator_names:
    if st.button("Submit for AI Evaluation"):
        
        final_transcript_to_evaluate = ""

        # ==========================================
        # PATH A: WE HAVE AUDIO (Needs AssemblyAI)
        # ==========================================
        if audio_bytes:
            if not aai_api_key:
                st.error("⚠️ Please enter your AssemblyAI API Key in the sidebar/secrets!")
                st.stop()
                
            with open("temp_audio.wav", "wb") as f:
                f.write(audio_bytes)
                
            with st.spinner(f"Transcribing audio for {candidate_name}..."):
                try:
                    loop = asyncio.get_event_loop()
                except RuntimeError:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    
                aai.settings.api_key = aai_api_key
                transcriber = aai.Transcriber()
                config = aai.TranscriptionConfig(
                    speaker_labels=True,
                    speakers_expected=total_speakers,
                    speech_models=["universal-2"] 
                )
                
                transcript = transcriber.transcribe("temp_audio.wav", config)

                if transcript.error:
                    st.error(f"Transcription failed: {transcript.error}")
                    st.stop()
                else:
                    st.success("Audio Transcription complete!")
                    
                    full_transcript_text = f"Interview Transcript: {candidate_name} & Evaluators: {evaluator_names}\n"
                    full_transcript_text += "="*50 + "\n\n"
                    for utterance in transcript.utterances:
                        speaker_name = f"Speaker {utterance.speaker}"
                        full_transcript_text += f"{speaker_name}: {utterance.text}\n\n"
                    
                    final_transcript_to_evaluate = full_transcript_text
                    
                    # Save Transcript to Dashboard
                    transcript_doc_name = f"{candidate_name}_Transcript.txt"
                    with open(os.path.join(SAVE_FOLDER, transcript_doc_name), "w") as f:
                        f.write(full_transcript_text)

        # ==========================================
        # PATH B: WE HAVE TEXT (Skip AssemblyAI)
        # ==========================================
        elif transcript_text:
            st.info("Direct text uploaded. Skipping AssemblyAI transcription.")
            
            final_transcript_to_evaluate = f"Interview Transcript: {candidate_name} & Evaluators: {evaluator_names}\n"
            final_transcript_to_evaluate += "="*50 + "\n\n"
            final_transcript_to_evaluate += transcript_text
            
            # Save Transcript to Dashboard
            transcript_doc_name = f"{candidate_name}_Transcript.txt"
            with open(os.path.join(SAVE_FOLDER, transcript_doc_name), "w", encoding="utf-8") as f:
                f.write(final_transcript_to_evaluate)

        # ==========================================
        # ANTHROPIC EVALUATION (Applies to both paths)
        # ==========================================
        if final_transcript_to_evaluate:
            if not anthropic_api_key:
                st.warning("Transcript ready! Add an Anthropic API key to your secrets to generate an evaluation.")
            else:
                with st.spinner("Claude is analyzing responses and generating the evaluation scorecard..."):
                    
                    client = Anthropic(api_key=anthropic_api_key)
                    
                    rubric_instructions = f"""
You are an expert MBA Admissions Committee Evaluator for a highly competitive program. Your task is to critically evaluate an incoming MBA candidate's interview transcript and generate a formal assessment report. The candidate's name is {candidate_name} and the evaluators were {evaluator_names}.

CRITICAL INSTRUCTIONS ON STRICTNESS & BIAS (TEMPERATURE CHECK):
You must act as a rigorous gatekeeper. Do not be lenient, polite, or forgiving in your internal assessment. Base your evaluation strictly on the empirical evidence present in the transcript. 
- Penalize vague answers, corporate jargon without substance, and failure to answer the specific question asked. However, Do not penalise for not knowing obscure facts; penalise for inability to reason from known facts. Further, a mere failure to recall a formula should not be penalized as heavily as conceptual error.
- Do not assume competence; the candidate must demonstrate it. However, candidates shall be scored only on what was actually tested and unasked questions shall not be charged to candidates.
- If a candidate struggles, do not sugarcoat the "Areas for Improvement."
- Ensure a fair, objective, and unflinching evaluation.
- No single parameter, however strong, can compensate for weakness in another. A candidate who scores A+ on Analytical Factor but C on Glocal Understanding does NOT qualify for an overall A grade. Each parameter is independently evaluated. The Overall Suitability is not an arithmetic average — it reflects the weakest-link principle.

GRADING SCALE:
Rate the candidate on the following 4 parameters using a scale from C- (Lowest) to A+ (Best). The valid grades are: A+, A, A-, B+, B, B-, C+, C, C-.
• C band (C-, C, C+): Below expectations. Significant gaps. Not ready without substantial development. Does not demonstrate the aptitude or mindset to benefit from or contribute to an MBA programme at this time.
• B band (B-, B, B+): Meets expectations with varying degrees of competence. Shows adequate aptitude and potential. Will grow with the right programme. Some gaps are expected and acceptable.
• A band (A-, A, A+): Exceeds expectations. Demonstrates clear thinking, maturity, and differentiated insight. A+ is reserved for truly exceptional responses that go beyond the question and demonstrate mastery. It should be rare.

EVALUATION PARAMETERS:
1. Glocal: 
Assess the candidate's understanding and awareness of current affairs and static knowledge. Evaluate their ability to analyze the implications of global/local events in given business or economic scenarios. However, we don’t expect candidates to know everything. Therefore, if candidate accepts he or she does not know answer to a particular questions but is able to answer other questions, do not over penalize.
Grade anchors: C band: Unaware of major issues. Cannot engage even when prompted with context. B band: Aware of some issues. Can discuss with some coherence. B+ shows a genuine or a well-reasoned opinion. A band: Goes beyond awareness to insight — connects issues, sees implications, shows genuine curiosity. A+ brings a perspective that is unexpected and substantiated.

2. Subject Knowledge: 
Based on the candidate's stated educational background, assess the depth and accuracy of their knowledge regarding the subjects they have formally studied. Look for clarity of fundamental concepts.
Grade anchors: C band: Struggles with basics of own domain. No real-world application visible. B band: Competent on core concepts. At least one real-world connection made. B+ shows genuine interest or reading beyond curriculum. A band: Strong command of their domain for their stage. Clear thinker. Can bridge theory to reality comfortably. A+ shows intellectual independence — a genuine point of view.

3. Work Experience:
- STRICT RULE: If the candidate has no formal, full-time work experience and has only completed internships, you MUST assign a score of "C-".
- If they have formal work experience: Judge them on the complexity of the work done and their depth of understanding of their domain/industry.
Grade Anchors: C band: Cannot describe any experience with specificity. No evidence of agency or reflection. B band: At least one experience described with reasonable specificity and a lesson extracted. B+ shows clear agency and some interpersonal or pressure-handling evidence. A band: Experiences — however modest in scale — are narrated with vivid specificity, clear agency, and genuine self-awareness. The candidate has clearly learned from what they have lived. A+ shows maturity or initiative that is exceptional for their age and background.

4. Analytical Skills: 
Gauge the candidate's ability to think analytically and their quantitative aptitude. Look for evidence of structured problem-solving, logical deduction, mental math capabilities, or how they approach puzzles and estimations.
grade anchors: C band: Refuses to engage or shows no structured thinking even after prompting. B band: Attempts most problems. Some structure visible. Reasonable on estimation. B+ handles the cross-stream question with genuine first-principles effort. A band: Clean, confident, structured — even on unfamiliar territory. Comfortable with uncertainty. A+ shows elegance and intellectual joy in problem solving.

OUTPUT FORMAT:
Generate the final report exactly in the structure below. Ensure the text is formatted cleanly.

Name of the candidate: {candidate_name}
Date of Interview: [Extract from transcript, or output "Not Provided"]

Executive Summary:
[Provide a concise, 3-4 sentence objective summary of the candidate's overall performance, highlighting their primary narrative and overall fit.]

Key Strengths:
* [Bullet point 1]
* [Bullet point 2]
* [Bullet point 3]

Areas for Improvement:
* [Bullet point 1]
* [Bullet point 2]
* [Bullet point 3]

Parameter 1 Score: [Grade]
Reason for the score in 2 lines: [Strictly 2 lines explaining the evidence for the Glocal score]

Parameter 2 Score: [Grade]
Reason for the score in 2 lines: [Strictly 2 lines explaining the evidence for the Subject Knowledge score]

Parameter 3 Score: [Grade]
Reason for the score in 2 lines: [Strictly 2 lines explaining the evidence for the Work Experience score. If C- due to internships, state this explicitly.]

Parameter 4 Score: [Grade]
Reason for the score in 2 lines: [Strictly 2 lines explaining the evidence for the Analytical Skills score]

Final Recommendation: 
[Strong Hire / Hire / Waitlist / Reject - followed by a 1-sentence justification]

OUTPUT CONSTRAINT: Do not use any Markdown formatting like asterisks (**) or hashes (#). Output strictly in plain text.
"""
                    
                    response = client.messages.create(
                        model="claude-haiku-4-5", 
                        max_tokens=1500, 
                        system=rubric_instructions, 
                        messages=[
                            {"role": "user", "content": f"Here is the interview transcript to evaluate:\n\n{final_transcript_to_evaluate}"}
                        ]
                    )
                    
                    evaluation_result = response.content[0].text
                    
                    st.subheader(f"📊 Evaluation Report: {candidate_name}")
                    st.markdown(evaluation_result)
                    
                    # Save to Word Document
                    doc = Document()
                    doc.add_heading(f'Evaluation Report: {candidate_name}', level=1)
                    
                    for line in evaluation_result.split('\n'):
                        doc.add_paragraph(line)
                        
                    eval_doc_name = f"{candidate_name}_Evaluation.docx"
                    doc_path = os.path.join(SAVE_FOLDER, eval_doc_name)
                    doc.save(doc_path)
                        
                    st.balloons()
